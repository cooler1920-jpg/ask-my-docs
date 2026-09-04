"""
File readers.

Every function here turns an uploaded file into plain text.
Images are described by Gemini Vision, then combined with whatever
explanation the user typed for that image.
"""

import csv
import io
import os

from google.genai import types

import rag_engine

TEXT_EXT = {".txt", ".md", ".markdown", ".log", ".json"}
IMAGE_EXT = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp"}

IMAGE_MIME = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".webp": "image/webp",
    ".gif": "image/gif",
    ".bmp": "image/bmp",
}

CAPTION_PROMPT = (
    "Describe this image thoroughly so it can be searched later. "
    "Read out ALL text visible in the image exactly as written, including numbers, "
    "names, dates, prices, headings and table contents. "
    "Then describe what the image shows. Be factual, do not speculate."
)


# --------------------------------------------------------------------------
# Chunking
# --------------------------------------------------------------------------

def chunk_text(text, size=900, overlap=150):
    """Split long text into overlapping pieces so context is never cut in half."""
    words = text.split()
    if not words:
        return []

    pieces = []
    step = max(size - overlap, 1)

    for start in range(0, len(words), step):
        piece = " ".join(words[start:start + size]).strip()
        if piece:
            pieces.append(piece)
        if start + size >= len(words):
            break

    return pieces


# --------------------------------------------------------------------------
# Individual formats
# --------------------------------------------------------------------------

PDF_OCR_PROMPT = (
    "Transcribe every word of this document, page by page. "
    "Preserve headings, numbering, tables and question numbers exactly as printed. "
    "Output plain text only. Do not summarise, do not add commentary."
)

MAX_INLINE_PDF = 18 * 1024 * 1024  # Gemini inline-data limit, with headroom


def read_pdf(data, client=None):
    """
    Read a PDF.

    Normal PDFs carry a text layer that pypdf extracts instantly and for free.
    Scanned PDFs (photographs of pages) have no text layer, so we hand the whole
    file to Gemini, which reads it the way a person would.
    """
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = []
    for number, page in enumerate(reader.pages, 1):
        text = (page.extract_text() or "").strip()
        if text:
            pages.append(f"[Page {number}]\n{text}")

    extracted = "\n\n".join(pages).strip()

    # A genuine text PDF yields plenty of characters per page.
    # Almost nothing means the pages are images.
    if len(extracted) >= 40 * max(len(reader.pages), 1):
        return extracted

    if client is None:
        return extracted

    if len(data) > MAX_INLINE_PDF:
        raise ValueError(
            "This looks like a scanned PDF and it is too large to read "
            f"({len(data) // (1024 * 1024)} MB). Please split it into smaller files."
        )

    response = client.models.generate_content(
        model=rag_engine.VISION_MODEL,
        contents=[
            types.Part.from_bytes(data=data, mime_type="application/pdf"),
            PDF_OCR_PROMPT,
        ],
    )
    return (response.text or "").strip()


def read_docx(data):
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def read_xlsx(data):
    from openpyxl import load_workbook

    workbook = load_workbook(io.BytesIO(data), data_only=True, read_only=True)
    parts = []

    for sheet in workbook.worksheets:
        parts.append(f"[Sheet: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            cells = ["" if v is None else str(v).strip() for v in row]
            if any(cells):
                parts.append(" | ".join(cells))

    workbook.close()
    return "\n".join(parts)


def read_csv(data):
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            text = data.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        text = data.decode("utf-8", errors="replace")

    rows = list(csv.reader(io.StringIO(text)))
    return "\n".join(" | ".join(cell.strip() for cell in row) for row in rows if any(row))


def read_txt(data):
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def read_image(client, data, extension, note=""):
    """Ask Gemini to describe the image, then add the user's own explanation."""
    mime = IMAGE_MIME.get(extension, "image/png")

    response = client.models.generate_content(
        model=rag_engine.VISION_MODEL,
        contents=[
            types.Part.from_bytes(data=data, mime_type=mime),
            CAPTION_PROMPT,
        ],
    )
    caption = (response.text or "").strip()

    if note.strip():
        return f"USER'S EXPLANATION OF THIS IMAGE:\n{note.strip()}\n\nWHAT THE IMAGE SHOWS:\n{caption}"
    return f"WHAT THE IMAGE SHOWS:\n{caption}"


# --------------------------------------------------------------------------
# Dispatcher
# --------------------------------------------------------------------------

def load(client, filename, data, note=""):
    """
    Turn one uploaded file into a list of {"text", "source"} chunks.
    Raises ValueError for unsupported file types.
    """
    extension = os.path.splitext(filename)[1].lower()

    if extension == ".pdf":
        text = read_pdf(data, client)
    elif extension == ".docx":
        text = read_docx(data)
    elif extension in {".xlsx", ".xlsm"}:
        text = read_xlsx(data)
    elif extension == ".csv":
        text = read_csv(data)
    elif extension in TEXT_EXT:
        text = read_txt(data)
    elif extension in IMAGE_EXT:
        text = read_image(client, data, extension, note)
    elif extension == ".doc":
        raise ValueError(
            "Old .doc files are not supported. Open it in Word and 'Save As' .docx."
        )
    elif extension == ".xls":
        raise ValueError(
            "Old .xls files are not supported. Open it in Excel and 'Save As' .xlsx."
        )
    else:
        raise ValueError(f"Unsupported file type: {extension or 'no extension'}")

    text = text.strip()
    if not text:
        raise ValueError(
            "No readable text found in this file. "
            "It may be empty, corrupted or password-protected."
        )

    # An image description is short and self-contained; keep it as one piece.
    if extension in IMAGE_EXT:
        return [{"text": text, "source": filename}]

    return [{"text": piece, "source": filename} for piece in chunk_text(text)]
